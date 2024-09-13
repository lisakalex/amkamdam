from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('English Tenses Table with Explanation', 0)

# Add the content of the table to the Word document
table_content = [
    ("TENSE", "FORM", "USE", "EXAMPLES"),
    ("PRESENT SIMPLE",
     "Affirmative: Subject + root form (додати 's' для he/she/it)\nNegative: Subject + do/does + not + root form\nQuestion: Do/Does + Subject + root form?",
     "Використовується для фактів, регулярних дій або звичок.",
     "1. She plays football every Sunday.\n2. I read books every evening.\n3. They live in London.\n4. He does not like coffee.\n5. Do you play tennis?"),

    ("PRESENT CONTINUOUS",
     "Affirmative: Subject + am/is/are + -ing form\nNegative: Subject + am/is/are + not + -ing form\nQuestion: Am/Is/Are + Subject + -ing form?",
     "Використовується для дій, які відбуваються прямо зараз або тимчасово.",
     "1. I am writing an email now.\n2. She is cooking dinner at the moment.\n3. They are playing football right now.\n4. I am not watching TV.\n5. Is he working?"),

    ("PAST SIMPLE",
     "Affirmative: Subject + past form\nNegative: Subject + did not + root form\nQuestion: Did + Subject + root form?",
     "Використовується для завершених дій у минулому.",
     "1. He visited Paris last year.\n2. I studied all night.\n3. They bought a new car.\n4. She did not go to the party.\n5. Did you see that movie?"),

    ("PAST CONTINUOUS",
     "Affirmative: Subject + was/were + -ing form\nNegative: Subject + was/were + not + -ing form\nQuestion: Was/Were + Subject + -ing form?",
     "Використовується для дій, які тривали в певний момент у минулому.",
     "1. I was studying when he called.\n2. They were sleeping at 10 PM.\n3. We were not playing football.\n4. Was she dancing when you arrived?"),

    ("PRESENT PERFECT",
     "Affirmative: Subject + have/has + past participle\nNegative: Subject + have/has + not + past participle\nQuestion: Have/Has + Subject + past participle?",
     "Використовується для дій, які відбулися в минулому і мають результат на теперішній момент.",
     "1. I have finished my homework.\n2. She has been to Japan.\n3. They have not met her yet.\n4. Have you seen this film before?"),

    ("PAST PERFECT",
     "Affirmative: Subject + had + past participle\nNegative: Subject + had not + past participle\nQuestion: Had + Subject + past participle?",
     "Використовується для дій, що завершилися до іншої дії в минулому.",
     "1. I had eaten before she arrived.\n2. They had left by the time I got there.\n3. She hadn't finished her homework.\n4. Had you finished by then?"),

    ("FUTURE SIMPLE",
     "Affirmative: Subject + will + root form\nNegative: Subject + will not (won’t) + root form\nQuestion: Will + Subject + root form?",
     "Використовується для дій у майбутньому або рішень, прийнятих у момент розмови.",
     "1. I will call you later.\n2. She will come tomorrow.\n3. They won’t go to the party.\n4. Will he help us?"),

    ("FUTURE CONTINUOUS",
     "Affirmative: Subject + will be + -ing form\nNegative: Subject + will not (won’t) be + -ing form\nQuestion: Will + Subject + be + -ing form?",
     "Використовується для дій, що триватимуть у певний момент у майбутньому.",
     "1. I will be working at 8 PM tomorrow.\n2. They won’t be waiting long.\n3. Will she be sleeping when we arrive?"),

    ("FUTURE PERFECT",
     "Affirmative: Subject + will have + past participle\nNegative: Subject + will not (won’t) have + past participle\nQuestion: Will + Subject + have + past participle?",
     "Використовується для дій, що завершаться до певного моменту в майбутньому.",
     "1. I will have finished my project by 6 PM.\n2. They won’t have arrived by 10 PM.\n3. Will you have completed your task by tomorrow?")
]

# Add each row of the table to the document
for row in table_content:
    doc.add_heading(row[0], level=1)  # Tense
    doc.add_paragraph(f"Form: {row[1]}\nUse: {row[2]}\nExamples:\n{row[3]}")

# Save the document
file_path_word = "English_Tenses_Table.docx"
doc.save(file_path_word)

# file_path_word
